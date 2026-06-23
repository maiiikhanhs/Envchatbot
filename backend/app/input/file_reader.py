from __future__ import annotations
"""Document reading for preprocessing before ComfyUI workflow execution."""

from pathlib import Path
from typing import Tuple, Dict

from docx import Document
from pypdf import PdfReader

from app.models.schemas import (
    SUPPORTED_FILE_EXTENSIONS,
    ERROR_FILE_NOT_FOUND,
    ERROR_UNSUPPORTED_FILE_TYPE,
    ERROR_FILE_TEXT_EXTRACTION_FAILED,
    ERROR_PDF_TEXT_NOT_FOUND,
)


class FileReaderError(Exception):
    """Base exception for file reader errors."""


class FileNotFoundCustomError(FileReaderError):
    """Raised when file does not exist."""


class UnsupportedFileTypeError(FileReaderError):
    """Raised when file type is not supported."""


class FileExtractionError(FileReaderError):
    """Raised when file content cannot be extracted."""


def get_file_info(file_path: str) -> Dict:
    """Return basic file metadata for downstream preprocessing steps."""
    path = Path(file_path)

    return {
        "file_name": path.name,
        "file_path": str(path),
        "file_extension": path.suffix.lower(),
        "file_size_bytes": path.stat().st_size,
    }


def extract_docx_text(file_path: str) -> str:
    """Extract text from a DOCX file for chunking."""
    try:
        document = Document(file_path)
        paragraphs = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        for table in document.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                cleaned_cells = [cell for cell in row_cells if cell]
                if cleaned_cells:
                    paragraphs.append(" | ".join(cleaned_cells))

        extracted_text = "\n".join(paragraphs).strip()

        if not extracted_text:
            raise FileExtractionError(ERROR_FILE_TEXT_EXTRACTION_FAILED)

        return extracted_text
    except FileReaderError:
        raise
    except Exception as exc:
        raise FileExtractionError(
            f"{ERROR_FILE_TEXT_EXTRACTION_FAILED}: {str(exc)}"
        ) from exc


def extract_pdf_text(file_path: str) -> str:
    """Extract text from a text-based PDF for chunking."""
    try:
        reader = PdfReader(file_path)
        pages_text = []

        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                pages_text.append(text.strip())

        extracted_text = "\n\n".join(pages_text).strip()

        if not extracted_text:
            raise FileExtractionError(ERROR_PDF_TEXT_NOT_FOUND)

        return extracted_text
    except FileReaderError:
        raise
    except Exception as exc:
        raise FileExtractionError(
            f"{ERROR_FILE_TEXT_EXTRACTION_FAILED}: {str(exc)}"
        ) from exc


def extract_file_content(file_path: str) -> Tuple[Dict, str]:
    """Read a supported file and return file_info plus extracted content."""
    path = Path(file_path)

    if not path.exists() or not path.is_file():
        raise FileNotFoundCustomError(ERROR_FILE_NOT_FOUND)

    extension = path.suffix.lower()

    if extension not in SUPPORTED_FILE_EXTENSIONS:
        raise UnsupportedFileTypeError(ERROR_UNSUPPORTED_FILE_TYPE)

    file_info = get_file_info(file_path)

    if extension == ".docx":
        file_content = extract_docx_text(file_path)
    elif extension == ".pdf":
        file_content = extract_pdf_text(file_path)
    else:
        raise UnsupportedFileTypeError(ERROR_UNSUPPORTED_FILE_TYPE)

    return file_info, file_content
