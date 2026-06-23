from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.worksheet.table import Table, TableStyleInfo


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "ui_blind_test_20_cases"
INPUT_DIR = OUT_DIR / "input_docx"
KEY_PATH = OUT_DIR / "answer_key_ui_blind_test_20_cases.xlsx"
README_PATH = OUT_DIR / "README.txt"


CASES = [
    {
        "id": "TC-001",
        "domain": "Không khí xung quanh",
        "standard": "QCVN 05:2023/BTNMT",
        "context": "Điểm quan trắc KK-01, trung bình 24 giờ, ngày 2026-02-12.",
        "applicability": "Thông số cơ bản trong không khí xung quanh.",
        "rows": [["SO2", "Trung bình 24 giờ", 118, "µg/Nm3"]],
        "expected": "Đạt",
        "basis": "SO2 trung bình 24 giờ: 118 <= 125 µg/Nm3.",
    },
    {
        "id": "TC-002",
        "domain": "Không khí xung quanh",
        "standard": "QCVN 05:2023/BTNMT",
        "context": "Điểm quan trắc KK-02, trung bình 1 giờ, ngày 2026-02-12.",
        "applicability": "Thông số cơ bản trong không khí xung quanh.",
        "rows": [["NO2", "Trung bình 1 giờ", 215, "µg/Nm3"]],
        "expected": "Chưa đạt",
        "basis": "NO2 trung bình 1 giờ: 215 > 200 µg/Nm3.",
    },
    {
        "id": "TC-003",
        "domain": "Không khí xung quanh",
        "standard": "QCVN 05:2023/BTNMT",
        "context": "Điểm quan trắc KK-03, trung bình 24 giờ, ngày 2026-03-20.",
        "applicability": "Thông số PM2,5 áp dụng giá trị từ ngày 2026-01-01.",
        "rows": [["PM2,5", "Trung bình 24 giờ", 48, "µg/Nm3"]],
        "expected": "Chưa đạt",
        "basis": "PM2,5 trung bình 24 giờ: 48 > 45 µg/Nm3.",
    },
    {
        "id": "TC-004",
        "domain": "Không khí xung quanh",
        "standard": "QCVN 05:2023/BTNMT",
        "context": "Điểm quan trắc KK-04, trung bình 8 giờ, ngày 2026-02-13.",
        "applicability": "Thông số cơ bản trong không khí xung quanh.",
        "rows": [["CO", "Trung bình 8 giờ", 9800, "µg/Nm3"]],
        "expected": "Đạt",
        "basis": "CO trung bình 8 giờ: 9.800 <= 10.000 µg/Nm3.",
    },
    {
        "id": "TC-005",
        "domain": "Không khí xung quanh",
        "standard": "QCVN 05:2023/BTNMT",
        "context": "Điểm quan trắc KK-05, trung bình năm 2026.",
        "applicability": "Thông số cơ bản trong không khí xung quanh.",
        "rows": [["Tổng bụi lơ lửng (TSP)", "Trung bình năm", 98, "µg/Nm3"]],
        "expected": "Đạt",
        "basis": "TSP trung bình năm: 98 <= 100 µg/Nm3.",
    },
    {
        "id": "TC-006",
        "domain": "Tiếng ồn",
        "standard": "QCVN 26:2010/BTNMT",
        "context": "Vị trí đo ON-01, khu vực đặc biệt, thời điểm 14:00.",
        "applicability": "Khung giờ từ 06:00 đến 21:00.",
        "rows": [["LAeq", "Khu vực đặc biệt / 14:00", 54, "dBA"]],
        "expected": "Đạt",
        "basis": "Khu vực đặc biệt ban ngày: 54 <= 55 dBA.",
    },
    {
        "id": "TC-007",
        "domain": "Tiếng ồn",
        "standard": "QCVN 26:2010/BTNMT",
        "context": "Vị trí đo ON-02, khu vực đặc biệt, thời điểm 22:30.",
        "applicability": "Khung giờ từ 21:00 đến 06:00.",
        "rows": [["LAeq", "Khu vực đặc biệt / 22:30", 47, "dBA"]],
        "expected": "Chưa đạt",
        "basis": "Khu vực đặc biệt ban đêm: 47 > 45 dBA.",
    },
    {
        "id": "TC-008",
        "domain": "Tiếng ồn",
        "standard": "QCVN 26:2010/BTNMT",
        "context": "Vị trí đo ON-03, khu vực thông thường, thời điểm 20:30.",
        "applicability": "Khung giờ từ 06:00 đến 21:00.",
        "rows": [["LAeq", "Khu vực thông thường / 20:30", 70, "dBA"]],
        "expected": "Đạt",
        "basis": "Khu vực thông thường ban ngày: 70 <= 70 dBA.",
    },
    {
        "id": "TC-009",
        "domain": "Tiếng ồn",
        "standard": "QCVN 26:2010/BTNMT",
        "context": "Vị trí đo ON-04, khu vực thông thường, thời điểm 05:30.",
        "applicability": "Khung giờ từ 21:00 đến 06:00.",
        "rows": [["LAeq", "Khu vực thông thường / 05:30", 56, "dBA"]],
        "expected": "Chưa đạt",
        "basis": "Khu vực thông thường ban đêm: 56 > 55 dBA.",
    },
    {
        "id": "TC-010",
        "domain": "Nước thải công nghiệp",
        "standard": "QCVN 40:2011/BTNMT",
        "context": "Mẫu NT-CN-01, cột A, Kq = 1, Kf = 1.",
        "applicability": "Nguồn tiếp nhận áp dụng cột A.",
        "rows": [["BOD5 (20°C)", "Mẫu đơn", 29, "mg/L"], ["COD", "Mẫu đơn", 70, "mg/L"]],
        "expected": "Đạt",
        "basis": "BOD5 29 <= 30 mg/L và COD 70 <= 75 mg/L theo cột A.",
    },
    {
        "id": "TC-011",
        "domain": "Nước thải công nghiệp",
        "standard": "QCVN 40:2011/BTNMT",
        "context": "Mẫu NT-CN-02, cột A, Kq = 1, Kf = 1.",
        "applicability": "Nguồn tiếp nhận áp dụng cột A.",
        "rows": [["Chất rắn lơ lửng", "Mẫu đơn", 55, "mg/L"]],
        "expected": "Chưa đạt",
        "basis": "Chất rắn lơ lửng 55 > 50 mg/L theo cột A.",
    },
    {
        "id": "TC-012",
        "domain": "Nước thải công nghiệp",
        "standard": "QCVN 40:2011/BTNMT",
        "context": "Mẫu NT-CN-03, cột B, Kq = 1, Kf = 1.",
        "applicability": "Nguồn tiếp nhận áp dụng cột B.",
        "rows": [["pH", "Mẫu đơn", 5.4, "-"]],
        "expected": "Chưa đạt",
        "basis": "pH 5,4 nằm dưới khoảng 5,5 đến 9 theo cột B.",
    },
    {
        "id": "TC-013",
        "domain": "Nước thải công nghiệp",
        "standard": "QCVN 40:2011/BTNMT",
        "context": "Mẫu NT-CN-04, cột B, Kq = 1, Kf = 1.",
        "applicability": "Nguồn tiếp nhận áp dụng cột B.",
        "rows": [["Amoni (tính theo N)", "Mẫu đơn", 10, "mg/L"], ["Tổng nitơ", "Mẫu đơn", 39, "mg/L"]],
        "expected": "Đạt",
        "basis": "Amoni 10 <= 10 mg/L và tổng nitơ 39 <= 40 mg/L theo cột B.",
    },
    {
        "id": "TC-014",
        "domain": "Nước thải công nghiệp",
        "standard": "QCVN 40:2011/BTNMT",
        "context": "Mẫu NT-CN-05, cột A, Kq = 1, Kf = 1.",
        "applicability": "Nguồn tiếp nhận áp dụng cột A.",
        "rows": [["Tổng dầu mỡ khoáng", "Mẫu đơn", 5.5, "mg/L"]],
        "expected": "Chưa đạt",
        "basis": "Tổng dầu mỡ khoáng 5,5 > 5 mg/L theo cột A.",
    },
    {
        "id": "TC-015",
        "domain": "Nước thải sinh hoạt",
        "standard": "QCVN 14:2008/BTNMT",
        "context": "Mẫu NT-SH-01, cột A, hệ số K = 1.",
        "applicability": "Nguồn tiếp nhận áp dụng cột A.",
        "rows": [["BOD5 (20°C)", "Mẫu đơn", 31, "mg/L"]],
        "expected": "Chưa đạt",
        "basis": "BOD5 31 > 30 mg/L theo cột A khi K = 1.",
    },
    {
        "id": "TC-016",
        "domain": "Nước thải sinh hoạt",
        "standard": "QCVN 14:2008/BTNMT",
        "context": "Mẫu NT-SH-02, cột B, hệ số K = 1.",
        "applicability": "Nguồn tiếp nhận áp dụng cột B.",
        "rows": [["Tổng chất rắn lơ lửng (TSS)", "Mẫu đơn", 100, "mg/L"], ["Amoni (tính theo N)", "Mẫu đơn", 9.8, "mg/L"]],
        "expected": "Đạt",
        "basis": "TSS 100 <= 100 mg/L và amoni 9,8 <= 10 mg/L theo cột B khi K = 1.",
    },
    {
        "id": "TC-017",
        "domain": "Nước thải sinh hoạt",
        "standard": "QCVN 14:2008/BTNMT",
        "context": "Mẫu NT-SH-03, cột A, hệ số K = 1.",
        "applicability": "Nguồn tiếp nhận áp dụng cột A.",
        "rows": [["pH", "Mẫu đơn", 4.9, "-"]],
        "expected": "Chưa đạt",
        "basis": "pH 4,9 nằm dưới khoảng 5 đến 9.",
    },
    {
        "id": "TC-018",
        "domain": "Nước mặt",
        "standard": "QCVN 08:2023/BTNMT",
        "context": "Mẫu NM-01, loại hình sông/suối/kênh/mương/khe/rạch.",
        "applicability": "Mức phân loại mục tiêu: B.",
        "rows": [["pH", "Mẫu đơn", 7.2, "-"], ["BOD5", "Mẫu đơn", 5.8, "mg/L"], ["COD", "Mẫu đơn", 14, "mg/L"], ["DO", "Mẫu đơn", 5.1, "mg/L"]],
        "expected": "Đạt",
        "basis": "Các thông số pH, BOD5, COD, DO đều nằm trong điều kiện mức B.",
    },
    {
        "id": "TC-019",
        "domain": "Nước mặt",
        "standard": "QCVN 08:2023/BTNMT",
        "context": "Mẫu NM-02, loại hình sông/suối/kênh/mương/khe/rạch.",
        "applicability": "Mức phân loại mục tiêu: B.",
        "rows": [["pH", "Mẫu đơn", 7.4, "-"], ["BOD5", "Mẫu đơn", 6.2, "mg/L"], ["COD", "Mẫu đơn", 16, "mg/L"], ["DO", "Mẫu đơn", 4.5, "mg/L"]],
        "expected": "Chưa đạt",
        "basis": "BOD5 6,2 > 6 mg/L, COD 16 > 15 mg/L và DO 4,5 < 5,0 mg/L so với mức B.",
    },
    {
        "id": "TC-020",
        "domain": "Khí thải công nghiệp",
        "standard": "QCVN 19:2009/BTNMT",
        "context": "Mẫu KT-01, cột B, Kp = 1, Kv = 1.",
        "applicability": "Thông số khí thải công nghiệp đối với bụi và chất vô cơ.",
        "rows": [["Lưu huỳnh đioxit (SO2)", "Mẫu khí thải", 520, "mg/Nm3"]],
        "expected": "Chưa đạt",
        "basis": "SO2 520 > 500 mg/Nm3 theo cột B khi Kp = 1 và Kv = 1.",
    },
]


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_grid(table, widths_dxa: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if style_name != "Heading 1" else 16)
        style.paragraph_format.space_after = Pt(6 if style_name != "Heading 1" else 8)


def add_metadata_table(doc: Document, case: dict) -> None:
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [2700, 6660]
    labels = [
        ("Mã hồ sơ", case["id"]),
        ("Nhóm dữ liệu", case["domain"]),
        ("Quy chuẩn tham chiếu", case["standard"]),
        ("Phạm vi áp dụng", case["applicability"]),
    ]
    for row, (label, value) in zip(table.rows, labels):
        set_cell_text(row.cells[0], label, bold=True, color="1F4D78")
        set_cell_text(row.cells[1], value)
        shade_cell(row.cells[0], "E8EEF5")
    set_table_grid(table, widths)


def add_measurement_table(doc: Document, case: dict) -> None:
    rows = case["rows"]
    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Thông số", "Kỳ đo / mô tả mẫu", "Giá trị ghi nhận", "Đơn vị"]
    widths = [2600, 3360, 1900, 1500]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True, color="1F4D78")
        shade_cell(cell, "E8EEF5")
    for row, values in zip(table.rows[1:], rows):
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value)
    set_table_grid(table, widths)


def build_docx(case: dict) -> Path:
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run(f"Hồ sơ quan trắc {case['id']}")
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    srun = subtitle.add_run("Tài liệu đầu vào cho kiểm thử UI")
    srun.font.name = "Calibri"
    srun.font.size = Pt(11)
    srun.font.color.rgb = RGBColor.from_string("555555")

    doc.add_heading("Thông tin hồ sơ", level=1)
    add_metadata_table(doc, case)

    doc.add_heading("Bối cảnh lấy mẫu", level=1)
    p = doc.add_paragraph(case["context"])
    p.paragraph_format.space_after = Pt(8)

    doc.add_heading("Số liệu ghi nhận", level=1)
    add_measurement_table(doc, case)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(4)
    nrun = note.add_run("Yêu cầu xử lý: trích xuất số liệu, xác định quy chuẩn liên quan và trả về nhận định theo quy tắc đối chiếu của hệ thống.")
    nrun.font.name = "Calibri"
    nrun.font.size = Pt(10)
    nrun.font.color.rgb = RGBColor.from_string("555555")

    doc.core_properties.title = f"Ho so quan trac {case['id']}"
    doc.core_properties.subject = "Bo du lieu dau vao kiem thu UI"
    doc.core_properties.comments = ""
    doc.core_properties.keywords = case["id"]

    path = INPUT_DIR / f"{case['id']}_ho_so_quan_trac.docx"
    doc.save(path)
    return path


def build_answer_key() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "answer_key"
    headers = [
        "case_id",
        "input_file",
        "domain",
        "standard",
        "expected_result",
        "reason",
        "do_not_include_in_ui_input",
    ]
    ws.append(headers)
    for case in CASES:
        ws.append(
            [
                case["id"],
                f"input_docx/{case['id']}_ho_so_quan_trac.docx",
                case["domain"],
                case["standard"],
                case["expected"],
                case["basis"],
                "TRUE",
            ]
        )

    header_fill = PatternFill("solid", fgColor="1F4D78")
    for cell in ws[1]:
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    widths = [12, 36, 26, 24, 16, 78, 24]
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[chr(64 + idx)].width = width
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:G{len(CASES) + 1}"
    table = Table(displayName="AnswerKey", ref=f"A1:G{len(CASES) + 1}")
    table.tableStyleInfo = TableStyleInfo(name="TableStyleMedium2", showRowStripes=True, showColumnStripes=False)
    ws.add_table(table)
    wb.save(KEY_PATH)


def build_readme() -> None:
    README_PATH.write_text(
        "\n".join(
            [
                "Bộ dữ liệu blind test UI - 20 case",
                "",
                "Cách dùng:",
                "1. Chỉ upload các file trong thư mục input_docx vào UI.",
                "2. Không upload file answer_key_ui_blind_test_20_cases.xlsx vào hệ thống cần kiểm thử.",
                "3. Sau khi UI/AI trả kết quả, đối chiếu theo cột expected_result và reason trong file đáp án.",
                "",
                "Nguyên tắc tạo file:",
                "- File input không có tiêu đề hoặc nội dung ghi sẵn kết quả đạt/chưa đạt.",
                "- Tên file chỉ chứa mã case trung lập.",
                "- Đáp án và ngưỡng đối chiếu nằm riêng trong workbook answer key.",
            ]
        ),
        encoding="utf-8",
    )


def verify_no_answer_leak() -> list[str]:
    leaks: list[str] = []
    banned = ["đạt", "chưa đạt", "không đạt", "khong dat", "dat"]
    for path in sorted(INPUT_DIR.glob("*.docx")):
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + " ".join(cell.text for cell in row.cells)
        lowered = text.lower()
        for word in banned:
            if word in lowered:
                leaks.append(f"{path.name}: {word}")
    return leaks


def main() -> None:
    if OUT_DIR.exists():
        shutil.rmtree(OUT_DIR)
    INPUT_DIR.mkdir(parents=True, exist_ok=True)

    for case in CASES:
        build_docx(case)
    build_answer_key()
    build_readme()

    leaks = verify_no_answer_leak()
    if leaks:
        raise RuntimeError("Answer leak detected: " + "; ".join(leaks))

    print(f"Created {len(CASES)} input DOCX files")
    print(KEY_PATH)
    print(README_PATH)


if __name__ == "__main__":
    main()
